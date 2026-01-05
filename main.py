import os
import requests
import time
import re
import json
from bs4 import BeautifulSoup  # NEW: Import BeautifulSoup for robust HTML parsing
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from threat_hunting_queries_recommendation_module import *

# --- CONFIGURATION ---
# IMPORTANT: Replace 'YOUR_API_KEY_HERE' with your actual Gemini API Key.
# It is highly recommended to use environment variables for keys in production.
API_KEY = os.getenv("GEMINI_API_KEY")

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", API_KEY)

# Change the model name from 'flash' to 'pro'
#API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"

INSTRUCTIONS_FILE = "instructions_threat_analysis.docx"
HTML_INPUT_FILE = r"Amazon Threat Intelligence identifies Russian cyber threat group targeting Western critical infrastructure _ AWS Security Blog.html"
ANALYSIS_OUTPUT_DOCX = "threat analysis_output.docx"
ANALYSIS_OUTPUT_MD = "threat analysis_output.md"

"""
def html_to_plain_text(html_content):

    #Performs basic conversion of HTML content to readable plain text.
    #Removes scripts, styles, and cleans up remaining tags.
    
    # 1. Remove script and style tags completely
    html_content = re.sub(r'<script.*?</script>', '', html_content, flags=re.IGNORECASE | re.DOTALL)
    html_content = re.sub(r'<style.*?</style>', '', html_content, flags=re.IGNORECASE | re.DOTALL)

    # 2. Replace common block elements with newlines for structure
    html_content = re.sub(r'</?(p|div|li|h[1-6]|br|blockquote)[^>]*>', '\n', html_content, flags=re.IGNORECASE)

    # 3. Remove remaining HTML tags (e.g., span, a, strong)
    text = re.sub(r'<[^>]+>', '', html_content)

    # 4. Clean up multiple newlines and excessive whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text).strip()

    return text
"""
# --- HELPER FUNCTIONS ---

def read_file_content(filepath):
    """Reads content from a specified file (.txt or .docx based on extension)."""

    if filepath.endswith('.docx'):
        try:
            document = Document(filepath)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except FileNotFoundError:
            print(f"ERROR: DOCX file not found at {filepath}")
            return None
        except PackageNotFoundError:
            print(f"ERROR: DOCX file at {filepath} is invalid or corrupt.")
            return None
        except Exception as e:
            print(f"ERROR reading DOCX file {filepath}: {e}")
            return None

    else:
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except FileNotFoundError:
            print(f"ERROR: File not found at {filepath}")
            return None
        except Exception as e:
            print(f"ERROR reading file {filepath}: {e}")
            return None


def html_to_plain_text(html_content):

    #Parses HTML content using BeautifulSoup to extract the main article text,
    #falling back to extracting text from all <p> tags if no main container is found.
 
    try:
        # Initialize BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Prioritize main content containers to avoid sidebars/headers
        article_tags = soup.find_all(
            ['article', 'div'],
            attrs={'id': re.compile(r'content|main', re.IGNORECASE),
                   'class': re.compile(r'article|main|post', re.IGNORECASE)}
        )

        article_text = ''

        if article_tags:
            # Extract text from the main article containers, using newlines for separation
            for tag in article_tags:
                article_text += tag.get_text(separator='\n\n', strip=True) + '\n\n'
        else:
            # Fallback: grab all <p> tags
            article_text = '\n\n'.join(p.get_text(strip=True) for p in soup.find_all('p'))

        # Clean up excessive whitespace that might result from the separators
        plain_text = re.sub(r'\n\s*\n', '\n\n', article_text).strip()

        if not plain_text:
            # Final fallback: if no paragraphs were found, just extract all visible text
            return soup.get_text(separator='\n', strip=True)

        return plain_text

    except Exception as e:
        print(f"ERROR during BeautifulSoup parsing: {e}")
        return html_content  # Return raw HTML as a safety fallback


def call_gemini_api(system_instruction, text_content, max_retries=5):
    """
    Calls the Gemini API with exponential backoff and enables Google Search grounding.
    """
    if GEMINI_API_KEY == "YOUR_API_KEY_HERE":
        print("ERROR: Please set your GEMINI_API_KEY in the script or environment variable.")
        return None

    # Adjusted query to reflect the threat intelligence task
    user_query = f"Perform the requested threat intelligence analysis on the following raw text data:\n\n---\n\n{text_content}"

    payload = {
        "contents": [{"parts": [{"text": user_query}]}],
        "systemInstruction": {
            "parts": [{"text": system_instruction}]
        },
        "tools": [{"google_search": {}}],
    }

    headers = {'Content-Type': 'application/json'}

    for attempt in range(max_retries):
        try:
            full_url = f"{API_URL}?key={GEMINI_API_KEY}"
            print(f"Attempting API call (Attempt {attempt + 1}/{max_retries})...")
            response = requests.post(full_url, headers=headers, data=json.dumps(payload))
            response.raise_for_status()

            result = response.json()
            generated_text = result.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text',
                                                                                                            'No content generated.')
            return generated_text

        except requests.exceptions.HTTPError as e:
            print(f"HTTP Error: {e.response.status_code} - {e.response.text}")
            if e.response.status_code in [429, 500, 503]:
                wait_time = 2 ** attempt
                print(f"Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                return f"API Failed (HTTP Error): {e.response.status_code}"

        except requests.exceptions.RequestException as e:
            print(f"An error occurred during the request: {e}")
            wait_time = 2 ** attempt
            print(f"Retrying in {wait_time} seconds...")
            time.sleep(wait_time)

    return "API call failed after maximum retries."


# --- DOCX PARSING & WRITING FUNCTIONS ---

def add_styled_paragraph(document, text):
    """Adds a paragraph, bolding text before the first colon/dash if it looks like a label."""
    text = text.strip()
    if ':' in text and not text.startswith(('*', '-')) and not re.match(r'^\d+\.\s', text):
        parts = text.split(':', 1)
        p = document.add_paragraph()
        p.add_run(parts[0].strip()).bold = True
        p.add_run(': ' + parts[1].strip())
    else:
        document.add_paragraph(text)


def parse_and_write_analysis(document, content):
    """
    Parses the structured analysis content and writes it to the DOCX document
    using appropriate formatting (Headings, Lists, Tables).
    """
    lines = content.strip().split('\n')
    in_table = False
    table_data = []

    for line in lines:
        stripped_line = line.strip()

        # --- TABLE PROCESSING ---
        if stripped_line.startswith(':---') or stripped_line.startswith('| :---'):
            in_table = True
            continue

        if in_table and stripped_line.startswith('|'):
            row_items = [item.strip().replace('**', '') for item in stripped_line.split('|')[1:-1]]
            table_data.append(row_items)
            continue

        if in_table and not stripped_line.startswith('|'):
            if table_data and len(table_data) > 0 and len(table_data[0]) > 0:
                try:
                    cols = len(table_data[0])
                    rows = len(table_data)
                    table = document.add_table(rows=rows, cols=cols)
                    table.style = 'Light Shading Accent 1'

                    for i, row_data in enumerate(table_data):
                        if len(row_data) == cols:
                            row_cells = table.rows[i].cells
                            for j, cell_text in enumerate(row_data):
                                cell = row_cells[j]
                                cell.text = cell_text
                                if i == 0:
                                    cell.paragraphs[0].runs[0].bold = True
                    document.add_paragraph()
                except Exception as e:
                    document.add_paragraph(f"--- Table Error: {e} ---")

            in_table = False
            table_data = []

        # --- NON-TABLE PROCESSING ---
        if re.match(r'^\d+\.\s+([A-Z].+)', stripped_line):
            heading_text = stripped_line.split('.', 1)[1].strip()
            document.add_heading(heading_text, level=1)
            document.add_paragraph()
            continue

        if stripped_line.startswith(('* ', '- ')):
            document.add_paragraph(stripped_line[2:], style='List Bullet')
        elif stripped_line:
            add_styled_paragraph(document, stripped_line)


def write_to_docx(filepath, content):
    """Initializes the DOCX document and calls the parser."""
    try:
        document = Document()
        document.add_heading('Gemini CTI Analysis Report', level=0)

        parse_and_write_analysis(document, content)

        document.save(filepath)
        print(f"\n[SUCCESS] Analysis output written and formatted to: {filepath}")
    except Exception as e:
        print(f"ERROR writing to DOCX file {filepath}: {e}")


# --- MARKDOWN WRITING FUNCTION ---

def write_to_markdown(filepath, content):
    """Writes the raw analysis string content directly to a Markdown (.md) file."""
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"\n[SUCCESS] Raw analysis output written to Markdown file: {filepath}")
    except Exception as e:
        print(f"ERROR writing to Markdown file {filepath}: {e}")


# --- MAIN EXECUTION ---

def main():
    print("--- Gemini LLM CTI Tool ---")

    # 1. Load System Instructions
    system_instruction = read_file_content(INSTRUCTIONS_FILE)
    if not system_instruction:
        return

    print(f"\n[SUCCESS] Loaded System Instructions from {INSTRUCTIONS_FILE}.")

    # 2. Load and Convert HTML Content
    html_content = read_file_content(HTML_INPUT_FILE)
    if not html_content:
        return

    plain_text_content = html_to_plain_text(html_content)

    #write_to_docx("plain_text_from_html.docx", plain_text_content)

    print(f"\n[SUCCESS] Loaded and Converted HTML from {HTML_INPUT_FILE}.")

    # 3. Call the Gemini API
    print("\nStarting analysis via Gemini API...")
    analysis_result = call_gemini_api(system_instruction, plain_text_content)

    # 4. Write Results to DOCX and Markdown
    if analysis_result and not analysis_result.startswith("API"):
        # Write the beautifully formatted DOCX file
        write_to_docx(ANALYSIS_OUTPUT_DOCX, analysis_result)

        # Write the portable Markdown file
        write_to_markdown(ANALYSIS_OUTPUT_MD, analysis_result)

    else:
        # If API failed, print the error message
        print("\n==============================================")
        print("           GEMINI ANALYSIS ERROR")
        print("==============================================")
        print(analysis_result)
        print("==============================================")

    print("threat hunting queries recommendation generation......")

    threat_hunting_queries_recommendation(plain_text_content)


if __name__ == "__main__":
    main()